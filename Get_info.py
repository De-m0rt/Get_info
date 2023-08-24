import wmi
import openpyxl
from openpyxl import load_workbook
import os
import sys
from barcode import Code128
from barcode.writer import ImageWriter
import docx
from docx.shared import Mm
from colorama import init, Fore, Style
import subprocess
import time
import re

lib_path = os.path.abspath(os.path.join(__file__))  # определяем текущий путь к файлу
sys.path.append(lib_path)  # записываем текущий путь в библиотеку


def in_exel(dict_for_sheet):
    fn = 'Get_info.xlsx'
    try:  # пытаемся открыть файл, если не выходит то создаем свой, если выходит, работаем с тем что открыли.
        wb = load_workbook(fn)  # открываем экселевский файл
    except FileNotFoundError:
        wb = openpyxl.Workbook()
    ws = wb.active  # работаем по умолчанию в первом листе открывшегося файла
    info = []
    for v in dict_for_sheet.values():  # перебираем все значения словаря
        for i in v:
            info.append(i)
    ws.append(info)  # записываем что данные из словаря в строку, ключ и значения в разные ячейки
    wb.save(fn)
    wb.close()


def in_word(dict_to_code, number_of_pages):
    fn = "Bar_code.docx"
    if os.path.exists(fn):
        doc = docx.Document(fn)
    else:
        doc = docx.Document()
    # ========== настройки формата документа
    sections = doc.sections
    for section in sections:
        section.page_height = Mm(50)
        section.page_width = Mm(70)
        section.left_margin = Mm(2)
        section.right_margin = Mm(2)
        section.top_margin = Mm(1)
        section.bottom_margin = Mm(0)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)
    # ==============================================
    my_code = Code128(dict_to_code['SerialNumber'][0], writer=ImageWriter())#'SerialNumber' '  ' +
    my_code.save("Code_image")
    pm = len(doc.paragraphs)
    for k in range(number_of_pages):  # количество страниц в документе
        doc.add_picture("Code_image.png", width=Mm(65), height=Mm(25))  # добавляем картинку со штрих-кодом
        paragraphs = doc.paragraphs  # создаем элемент параграфа документа
        # (почему-то после добавления картинки нужно каждый раз это делать)
        for i in range(len(dict_to_code['MacAddress'])):  # добавляем строки с мак-адресами
            paragraphs[k + pm].add_run("Mac " + str(i + 1) + '  ' + dict_to_code['MacAddress'][i] + '\n')
        # paragraphs[k+pm].add_run("WinKey" + '  ' + dict_to_code['WinKey'][0] +'\n') #добавляем ключ венды
        paragraphs[k + pm].alignment = 1
    doc.save(fn)
    os.remove("Code_image.png")


def get_info():
    c = wmi.WMI()
    info = {}
    info.setdefault("CSName", [c.Win32_OperatingSystem()[0].CSName])  # имя компьютера на котором запущен процесс
    info.setdefault("SerialNumber", [c.Win32_BIOS()[0].SerialNumber])  # получаем серийник винды
    info.setdefault("WinKey", [c.Win32_OperatingSystem()[0].SerialNumber])  # ХЗ что это
    info.setdefault("OriginalProductKey", [c.SoftwareLicensingService()[0].OA3xOriginalProductKey])  # OEM Ключ ?
    for i in c.win32_physicalmedia():  # получаем серийники у тех дисков(HDD/CD/SD) у которых они есть
        if i.SerialNumber is not None:
            info["DiskSerial"] = info.get("DiskSerial", []) + [i.SerialNumber]
    for i in c.win32_physicalmemory():  # получаем серийники RAMM
        info["MemSerial"] = info.get("MemSerial", []) + [i.SerialNumber]
    for i in c.win32_networkadapterconfiguration():  # получаем все активные макадкреса
        if i.IPEnabled is True:
            info["MacAddress"] = info.get("MacAddress", []) + [i.MACAddress]
    return info


def wifi_connect(name, password, time_sleep):
    fn = name + ".xml"
    if not (os.path.exists(fn)):
        print("Wifi profile not found, create a new one")
        config = f"""<?xml version="1.0"?>
<WLANProfile xmlns="http://www.microsoft.com/networking/WLAN/profile/v1">
	<name>{name}</name>
	<SSIDConfig>
		<SSID>
			<name>{name}</name>
		</SSID>
	</SSIDConfig>
	<connectionType>ESS</connectionType>
	<connectionMode>auto</connectionMode>
	<MSM>
		<security>
			<authEncryption>
				<authentication>WPA2PSK</authentication>
				<encryption>AES</encryption>
				<useOneX>false</useOneX>
			</authEncryption>
			<sharedKey>
				<keyType>passPhrase</keyType>
				<protected>false</protected>
				<keyMaterial>{password}</keyMaterial>
			</sharedKey>
		</security>
	</MSM>
	<MacRandomization xmlns="http://www.microsoft.com/networking/WLAN/profile/v3">
		<enableRandomization>false</enableRandomization>
	</MacRandomization>
</WLANProfile>"""
        with open(fn, 'w') as file:
            file.write(config)
    else:
        print("Wifi profile loaded")
    os.system(f"netsh wlan add profile filename={fn}")
    print("Connecting wifi...")
    os.system(f"netsh wlan connect name={name} ssid ={name}")
    print(f"Wifi name = {Fore.YELLOW}{name}{Style.RESET_ALL}")
    print(f"Wifi pass = {Fore.YELLOW}{password}{Style.RESET_ALL}")

    time.sleep(time_sleep)
    try:
        wifi = subprocess.check_output(['netsh', 'WLAN', 'show', 'interfaces'])
    except subprocess.CalledProcessError:
        wifi = ""
    if name in str(wifi):
        print(f"{Fore.GREEN}Wifi connected to {name}{Style.RESET_ALL}")
    else:
        print(f"{Fore.RED}Can't connected to {name}{Style.RESET_ALL}")


def usr_options():
    fn = "Get_info_Options.txt"
    try:
        with open(fn) as f:
            s = f.read()
    except:
        with open(fn, "a") as f:
            s = f.write("""Number of pages (Bar_code.docx) = 1
Wifi_name = None
Wifi_pass = None
Sleep_time = 3""")
        f.close()
        with open(fn) as f:
            s = f.read()
    usr_conf = []
    pattern = ".+=\D*(\d+)\n*.+= *(.*)\n*.+= *(.*)\n*.*=.*(\d)"
    match = re.findall(pattern, s, re.MULTILINE)
    for i in match[0]:
        usr_conf.append(i)
    return usr_conf


def main():
    conf = usr_options()
    wifi_connect(conf[1], conf[2], int(conf[3]))
    if os.path.isfile("Get_info.xlsx"):
        print(
            f"{Fore.GREEN}Get_info.xlsx will be added{Style.RESET_ALL}")
    else:
        print(f"{Fore.RED}Get_info.xlsx not found, create a new one {Style.RESET_ALL}")

    if os.path.isfile("Bar_code.docx"):
        print(f"{Fore.GREEN}Bar_code.docx will be added{Style.RESET_ALL}")
    else:
        print(f"{Fore.RED}Bar_code.docx not found, create a new one{Style.RESET_ALL}")
    init()
    info = get_info()
    in_exel(info)
    in_word(info, int(conf[0]))
    for k, v in info.items():
        f = ''
        for i in v:
            f += (Fore.RED + "NO ITEM !" + Style.RESET_ALL) if i == '' else i + "   "
        print(k + " = " + f)
    os.system('pause')


if __name__ == main():
    main()
